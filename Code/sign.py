import tkinter as tk
from docx import Document


def sign():
    docx = Document()
    docx.add_paragraph(et_cn.get(), style='List Bullet')
    docx.add_paragraph(et_en.get(), style='List Bullet')
    docx.add_paragraph(et_mobile.get(), style='List Bullet')
    docx.add_paragraph(et_tel.get(), style='List Bullet')
    docx.add_paragraph(et_mail.get(), style='List Bullet')
    docx.save('signtest.docx')

    print(et_cn.get())

'''
GUI
'''
window = tk.Tk()
window.title('sign')
window.geometry('800x600')

'''
Label
'''
lb_title = tk.Label(window, text='sign', font=('Time new roman', 54))
lb_title.pack()
lb_cn = tk.Label(window, text='Name(ZH)', font=('Time new roman', 36))
lb_cn.place(x=0, y=100)
lb_en = tk.Label(window, text='Name(EN)', font=('Time new roman', 36))
lb_en.place(x=0, y=200)
lb_mobile = tk.Label(window, text='Mobile', font=('Time new roman', 36))
lb_mobile.place(x=0, y=300)
lb_tel = tk.Label(window, text='Tel', font=('Time new roman', 36))
lb_tel.place(x=0, y=400)
lb_mail = tk.Label(window, text='E-mail', font=('Time new roman', 36))
lb_mail.place(x=0, y=500)

'''
Entry
'''
et_cn = tk.Entry(window, font=('Time new roman', 36))
et_cn.place(x=200, y=100)
et_en = tk.Entry(window, font=('Time new roman', 36))
et_en.place(x=200, y=200)
et_mobile = tk.Entry(window, font=('Time new roman', 36))
et_mobile.place(x=200, y=300)
et_tel = tk.Entry(window, font=('Time new roman', 36))
et_tel.place(x=200, y=400)
et_mail = tk.Entry(window, font=('Time new roman', 36))
et_mail.place(x=200, y=500)


'''
Button
'''
bt_sign = tk.Button(window, text='sign', command=sign, font=('Time new roman', 36))
bt_sign.place(x=0, y=550)


window.mainloop()
